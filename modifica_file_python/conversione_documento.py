from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
import os

# Paths
md_path = "documento_originale_in.md"
img_path = "/image.png"
output_path = "/nome_docu"

# Read markdown content
with open(md_path, "r", encoding="utf-8") as f:
    md_content = f.read()

# Create document
doc = Document()

# Title
title = doc.add_heading("Documentazione Tecnica — Moduli HMI Marini", level=1)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Add architecture image
if os.path.exists(img_path):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run()
    run.add_picture(img_path, width=Inches(6.5))

    caption = doc.add_paragraph("Figura 1 — Architettura generale dei moduli")
    caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    caption.runs[0].italic = True

# Add markdown content as plain text paragraphs
for line in md_content.splitlines():
    if line.strip().startswith("# "):
        doc.add_heading(line.replace("# ", "").strip(), level=1)
    elif line.strip().startswith("## "):
        doc.add_heading(line.replace("## ", "").strip(), level=2)
    elif line.strip().startswith("### "):
        doc.add_heading(line.replace("### ", "").strip(), level=3)
    else:
        p = doc.add_paragraph(line)
        p.style.font.size = Pt(10)

# Save document
doc.save(output_path)

print(f"Documento creato: {output_path}")
